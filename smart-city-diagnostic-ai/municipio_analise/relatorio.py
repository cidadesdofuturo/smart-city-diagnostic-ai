from __future__ import annotations

from pathlib import Path

from docx import Document

from .utils import nome_arquivo_seguro, separar_paragrafos


NOMES_EXIBICAO = {
    "economica": "Dimensão Econômica",
    "sociocultural": "Dimensão Sociocultural",
    "meio_ambiente": "Dimensão Meio Ambiente",
    "capacidades_institucionais": "Dimensão Capacidades Institucionais",
}


def gerar_relatorio_word(municipio: str, relatorio: dict, output_dir: Path | str) -> Path:
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(f"Relatório Institucional – {municipio}", level=1)

    doc.add_heading("Análise Geral", level=2)
    for paragrafo in separar_paragrafos(relatorio["analise_geral"]):
        doc.add_paragraph(paragrafo)

    for chave, titulo in NOMES_EXIBICAO.items():
        bloco = relatorio["dimensoes"][chave]
        doc.add_heading(titulo, level=2)
        for paragrafo in separar_paragrafos(bloco["analise"]):
            doc.add_paragraph(paragrafo)
        doc.add_paragraph("Sugestões de melhoria:")
        for sugestao in bloco["sugestoes"]:
            doc.add_paragraph(sugestao, style="List Bullet")

    caminho = output_dir / f"Relatorio_{nome_arquivo_seguro(municipio)}.docx"
    doc.save(caminho)
    return caminho
