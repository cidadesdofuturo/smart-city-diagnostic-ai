# =====================================================
# GERAÇÃO DO RELATÓRIO WORD (adaptado à saída estruturada)
# =====================================================
from docx import Document

from . import config

NOMES_EXIBICAO = {
    "economica": "Dimensão Econômica",
    "sociocultural": "Dimensão Sociocultural",
    "meio_ambiente": "Dimensão Meio Ambiente",
    "capacidades_institucionais": "Dimensão Capacidades Institucionais",
}


def gerar_relatorio_word(municipio, relatorio: dict, pasta_saida=None):
    pasta_saida = pasta_saida or config.BASE_PATH
    doc = Document()
    doc.add_heading(f"Relatório Institucional – {municipio}", level=1)

    doc.add_heading("Análise Geral", level=2)
    doc.add_paragraph(relatorio["analise_geral"])

    for chave, titulo in NOMES_EXIBICAO.items():
        bloco = relatorio["dimensoes"][chave]
        doc.add_heading(titulo, level=2)
        doc.add_paragraph(bloco["analise"])
        doc.add_paragraph("Sugestões de melhoria:")
        for sugestao in bloco["sugestoes"]:
            doc.add_paragraph(sugestao, style="List Bullet")

    nome_arquivo = pasta_saida / f"Relatorio_{municipio.replace(' ', '_')}.docx"
    doc.save(nome_arquivo)
    print(f"📄 Relatório gerado: {nome_arquivo}")
    return nome_arquivo
